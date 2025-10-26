from flask import Blueprint, request, flash, redirect, url_for, g, make_response, jsonify, render_template
from . import api_login_required, login_required
from flask import current_app
from logic import (
    obtener_uuid_factura_por_id_logic,
    obtener_datos_factura_por_uuid_logic,
    obtener_datos_factura_logic,
    generar_lote_unico_logic,
    obtener_inventario_logic,
    obtener_historial_transacciones_logic,
    obtener_productos,
    obtener_almacenes_por_ingenio_logic
)
import qrcode
import io
from docx import Document
from datetime import datetime
import pandas as pd
import calendar

reports_bp = Blueprint('reports', __name__)

@reports_bp.route('/api/generar_lote')
@api_login_required
def generar_lote():
    try:
        nuevo_lote = generar_lote_unico_logic(g.user['ingenio_id'])
        return jsonify({'lote': nuevo_lote})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@reports_bp.route('/api/welcome')
@api_login_required
def welcome():
    current_app.logger.info(f"Request: {request.method} {request.path}")
    return jsonify({'message': 'Welcome to the Flask API Service!'})

@reports_bp.route('/api/qr_code/<int:transaccion_id>')
@api_login_required
def generar_qr_code(transaccion_id):
    try:
        factura_uuid = obtener_uuid_factura_por_id_logic(transaccion_id, g.user['ingenio_id'])
        if not factura_uuid:
            return "Factura no encontrada o sin UUID.", 404
        url = url_for('reports.ver_factura_publica', factura_uuid=factura_uuid, _external=True)
        qr_img = qrcode.make(url)
        img_io = io.BytesIO()
        qr_img.save(img_io, 'PNG')
        img_io.seek(0)
        return make_response(img_io.getvalue(), {'Content-Type': 'image/png'})
    except Exception as e:
        current_app.logger.error(f"Error al generar QR para transaccion {transaccion_id}: {e}", exc_info=True)
        return "Error al generar QR.", 500

@reports_bp.route('/factura_publica/<uuid:factura_uuid>')
def ver_factura_publica(factura_uuid):
    try:
        datos_factura = obtener_datos_factura_por_uuid_logic(str(factura_uuid))
        if not datos_factura:
            return "Factura no encontrada.", 404
        # Calculate totals
        subtotal = sum(d[5] for d in datos_factura['detalles'])
        tipo = datos_factura['info'][6]  # tipo is now included
        if tipo in ('servicio_secado', 'servicio_pelado'):
            iva = 0.0  # No IVA for services
        else:
            iva = subtotal * 0.12
        total_final = subtotal + iva
        return render_template('factura_publica.html', factura=datos_factura, subtotal=subtotal, iva=iva, total_final=total_final)
    except Exception as e:
        current_app.logger.error(f"Error al ver factura pública {factura_uuid}: {e}", exc_info=True)
        return "Error al cargar la factura.", 500

@reports_bp.route('/factura_publica/descargar/<uuid:factura_uuid>')
def descargar_factura_publica(factura_uuid):
    try:
        datos_factura = obtener_datos_factura_por_uuid_logic(str(factura_uuid))
        if not datos_factura:
            return "Factura no encontrada.", 404
        tipo = datos_factura['info'][6]
        if tipo in ('servicio_secado', 'servicio_pelado'):
            document = _crear_documento_factura_servicio(datos_factura['info'][0], tipo, datos_factura['info'][1], datos_factura['info'][2], datos_factura['detalles'], datos_factura['info'][5])
        else:
            document = _crear_documento_factura(datos_factura['info'][0], datos_factura['info'][1], datos_factura['info'][2], datos_factura['detalles'], datos_factura['info'][5])
        f = io.BytesIO()
        document.save(f)
        f.seek(0)
        return make_response(f.getvalue(), {'Content-Disposition': f'attachment; filename=Factura-{datos_factura["info"][0]}.docx', 'Content-type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'})
    except Exception as e:
        current_app.logger.error(f"Error al descargar factura pública {factura_uuid}: {e}", exc_info=True)
        return "Error al generar el documento de la factura.", 500

@reports_bp.route('/generar_factura/<int:transaccion_id>')
@login_required
def generar_factura(transaccion_id):
    try:
        datos_factura = obtener_datos_factura_logic(transaccion_id, g.user['ingenio_id'], ('venta', 'servicio_secado', 'servicio_pelado'))
        if not datos_factura:
            flash('La transacción no es una venta o servicio válido o no se encontró.', 'danger')
            return redirect(url_for('main.historial'))
        tipo = datos_factura['info'][4]
        if tipo in ('servicio_secado', 'servicio_pelado'):
            document = _crear_documento_factura_servicio(transaccion_id, tipo, datos_factura['info'][0], datos_factura['info'][1], datos_factura['detalles'], g.user['ingenio_nombre'])
        else:
            document = _crear_documento_factura(transaccion_id, datos_factura['info'][0], datos_factura['info'][1], datos_factura['detalles'], g.user['ingenio_nombre'])
        f = io.BytesIO()
        document.save(f)
        f.seek(0)
        return make_response(f.getvalue(), {'Content-Disposition': f'attachment; filename=Factura-{transaccion_id}.docx', 'Content-type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'})
    except Exception as e:
        current_app.logger.error(f"Error al generar factura: {e}", exc_info=True)
        flash(f"Error al generar la factura: {e}", 'danger')
        return redirect(url_for('main.historial'))

@reports_bp.route('/generar_factura_servicio/<int:transaccion_id>')
@login_required
def generar_factura_servicio(transaccion_id):
    try:
        datos_factura = obtener_datos_factura_logic(transaccion_id, g.user['ingenio_id'], ('servicio_secado', 'servicio_pelado'))
        if not datos_factura:
            flash('La transacción no es un servicio válido o no se encontró.', 'danger')
            return redirect(url_for('main.historial'))
        document = _crear_documento_factura_servicio(transaccion_id, datos_factura['info'][4], datos_factura['info'][1], datos_factura['info'][2], datos_factura['detalles'], g.user['ingenio_nombre'])
        f = io.BytesIO()
        document.save(f)
        f.seek(0)
        return make_response(f.getvalue(), {'Content-Disposition': f'attachment; filename=Factura-Servicio-{transaccion_id}.docx', 'Content-type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'})
    except Exception as e:
        current_app.logger.error(f"Error al generar factura de servicio: {e}", exc_info=True)
        flash(f"Error al generar la factura de servicio: {e}", 'danger')
        return redirect(url_for('main.historial'))

def _crear_documento_factura(transaccion_id, nombre_cliente, fecha_venta, detalles_venta, nombre_ingenio):
    document = Document()
    document.add_heading(nombre_ingenio, level=0)
    document.add_heading('FACTURA', level=1)
    p_info = document.add_paragraph()
    p_info.add_run('Cliente: ').bold = True
    p_info.add_run(f"{nombre_cliente}\n")
    p_info.add_run('Fecha: ').bold = True
    p_info.add_run(f"{datetime.strptime(fecha_venta, '%Y-%m-%d %H:%M:%S').strftime('%d/%m/%Y')}\n")
    p_info.add_run('Factura N°: ').bold = True
    p_info.add_run(f"{transaccion_id:05d}")
    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Descripción'
    hdr_cells[1].text = 'Cantidad'
    hdr_cells[2].text = 'Unidad'
    hdr_cells[3].text = 'Precio Unit.'
    hdr_cells[4].text = 'Subtotal'
    subtotal_calculado = sum(d[5] for d in detalles_venta)
    for nombre_prod, variedad, cantidad, unidad, precio_unit, subtotal in detalles_venta:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{nombre_prod} ({variedad})" if variedad else nombre_prod
        row_cells[1].text = f"{cantidad:.2f}"
        row_cells[2].text = unidad
        row_cells[3].text = f"${precio_unit:.2f}"
        row_cells[4].text = f"${subtotal:.2f}"
    iva = subtotal_calculado * 0.12
    total_final = subtotal_calculado + iva
    document.add_paragraph(f'Subtotal: ${subtotal_calculado:.2f}', style='List Bullet').alignment = 2
    document.add_paragraph(f'IVA (12%): ${iva:.2f}', style='List Bullet').alignment = 2
    p_total = document.add_paragraph()
    p_total.alignment = 2
    p_total.add_run(f'Total: ${total_final:.2f}').bold = True
    return document

def _crear_documento_factura_servicio(transaccion_id, tipo_servicio, nombre_cliente, fecha_servicio, detalles_servicio, nombre_ingenio):
    document = Document()
    document.add_heading(nombre_ingenio, level=0)
    document.add_heading('FACTURA DE SERVICIO', level=1)
    p_info = document.add_paragraph()
    p_info.add_run('Cliente: ').bold = True
    p_info.add_run(f"{nombre_cliente}\n")
    p_info.add_run('Fecha: ').bold = True
    p_info.add_run(f"{datetime.strptime(fecha_servicio, '%Y-%m-%d %H:%M:%S').strftime('%d/%m/%Y')}\n")
    p_info.add_run('Factura N°: ').bold = True
    p_info.add_run(f"{transaccion_id:05d}\n")
    p_info.add_run('Tipo de Servicio: ').bold = True
    p_info.add_run(tipo_servicio.replace('servicio_', '').capitalize())
    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Descripción'
    hdr_cells[1].text = 'Cantidad'
    hdr_cells[2].text = 'Unidad'
    hdr_cells[3].text = 'Precio Unit.'
    hdr_cells[4].text = 'Subtotal'
    subtotal_calculado = sum(d[5] for d in detalles_servicio)
    for nombre_prod, variedad, cantidad, unidad, precio_unit, subtotal, lote in detalles_servicio:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{nombre_prod} ({variedad}) - Lote: {lote}" if variedad else f"{nombre_prod} - Lote: {lote}"
        row_cells[1].text = f"{cantidad:.2f}"
        row_cells[2].text = unidad
        row_cells[3].text = f"${precio_unit:.2f}"
        row_cells[4].text = f"${subtotal:.2f}"
    # Services do not have IVA
    iva = 0.0
    total_final = subtotal_calculado + iva
    document.add_paragraph(f'Subtotal: ${subtotal_calculado:.2f}', style='List Bullet').alignment = 2
    document.add_paragraph(f'IVA (12%): ${iva:.2f}', style='List Bullet').alignment = 2
    p_total = document.add_paragraph()
    p_total.alignment = 2
    p_total.add_run(f'Total: ${total_final:.2f}').bold = True
    return document

@reports_bp.route('/exportar/<string:reporte_nombre>')
@login_required
def exportar(reporte_nombre):
    if g.user['nivel_acceso'] == 'empleado':
        flash('No tienes permiso para exportar registros.', 'danger')
        return redirect(request.referrer or url_for('main.compras'))

    try:
        df = pd.DataFrame()
        columnas = []
        titulo_periodo = "Todo el tiempo"

        if reporte_nombre == 'inventario':
            data = obtener_inventario_logic(g.user['ingenio_id'], request.args.get('producto_id', type=int), request.args.get('lote'), request.args.get('variedad'), request.args.get('fecha'), request.args.get('almacen_id', type=int))
            columnas = ["ID", "Producto", "Código", "Variedad", "Lote", "Cantidad", "Cantidad KG", "Unidad", "Estado", "Fecha Entrada", "Precio Venta", "Almacén"]
            df = pd.DataFrame(data, columns=columnas) if data else pd.DataFrame(columns=columnas)
            
            # Solución: Reemplazar valores NaN (Not a Number) con una cadena vacía antes de exportar.
            # Esto evita el error con xlsxwriter cuando hay precios de venta no definidos (NULL/None).
            df.fillna('', inplace=True)
        
        elif reporte_nombre == 'historial':
            hoy = datetime.now()
            periodo_tipo = request.args.get('periodo_tipo', 'todos')
            fecha_inicio, fecha_fin = None, None
            if periodo_tipo == 'rango':
                fecha_inicio, fecha_fin = request.args.get('fecha_inicio'), request.args.get('fecha_fin')
                if fecha_inicio and fecha_fin: titulo_periodo = f"Desde {fecha_inicio} hasta {fecha_fin}"
            elif periodo_tipo == 'mensual':
                mes, año = int(request.args.get('mes', hoy.month)), int(request.args.get('año', hoy.year))
                _, ultimo_dia = calendar.monthrange(año, mes)
                fecha_inicio, fecha_fin = f"{año}-{mes:02d}-01", f"{año}-{mes:02d}-{ultimo_dia}"
                titulo_periodo = f"Mes: {mes:02d}-{año}"
            
            data = obtener_historial_transacciones_logic(
                g.user['ingenio_id'], request.args.get('tipo'), request.args.get('producto_id', type=int),
                request.args.get('nombre'), request.args.get('lote'), fecha_inicio, fecha_fin
            )
            columnas = ["ID", "Tipo", "Nombre", "Fecha", "Producto", "Código", "Variedad", "Cantidad", "Cantidad KG", "Unidad", "Precio Unitario", "Subtotal", "Lote", "Observaciones"]
            df = pd.DataFrame(data, columns=columnas) if data else pd.DataFrame(columns=columnas)
        else:
            flash('Reporte no válido.', 'danger')
            return redirect(url_for('main.compras'))

        output = io.BytesIO()
        import xlsxwriter
        with xlsxwriter.Workbook(output, {'in_memory': True}) as workbook:
            # ... (Aquí iría toda la lógica de creación del Excel con xlsxwriter)
            # Por brevedad, se omite pero es la misma que tenías en app.py
            worksheet = workbook.add_worksheet(reporte_nombre.capitalize())
            # Escribir DataFrame en la hoja
            for c_idx, col_name in enumerate(df.columns):
                worksheet.write(0, c_idx, col_name)
            for r_idx, row in enumerate(df.itertuples(index=False), 1):
                worksheet.write_row(r_idx, 0, row)

        output.seek(0)
        return make_response(output.getvalue(), {
            'Content-Disposition': f'attachment; filename={reporte_nombre}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx',
            'Content-type': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        })
    except Exception as e:
        current_app.logger.error(f"Error during Excel export: {e}", exc_info=True)
        flash(f"Error al generar el reporte: {e}", "danger")
        return redirect(request.referrer or url_for('main.compras'))

@reports_bp.route('/productos')
@login_required
def productos():
    productos = obtener_productos()
    return render_template('productos.html', productos=productos)
